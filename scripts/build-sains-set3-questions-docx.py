import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "js" / "practice-science-sets.js"
OUTPUT = ROOT / "Senarai_Soalan_Bergambar_Sains_Set_3.docx"

MAROON = "8A1538"
PALE = "FCE7EF"
LIGHT = "FFF7FA"
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(100, 116, 139)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def load_questions():
    source = SOURCE.read_text(encoding="utf-8")
    match = re.search(r"const scienceSets = (\{[\s\S]*?\});\s+window\.PKSK_SET_QUESTIONS", source)
    if not match:
        raise RuntimeError("Data scienceSets tidak dijumpai.")
    science_sets = json.loads(match.group(1))
    questions = science_sets["3"]
    return [(index, item["q"]) for index, item in enumerate(questions, start=1) if item.get("fig")]


def build_document(questions):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("PKSK  |  LATIHAN SAINS")
    set_run_font(run, size=9, bold=True, color=RGBColor(138, 21, 56))

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("Senarai Soalan Bergambar")
    set_run_font(run, size=23, bold=True, color=RGBColor(138, 21, 56))

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Subjek Sains - Set 3")
    set_run_font(run, size=14, bold=True, color=INK)

    summary = doc.add_paragraph()
    summary.paragraph_format.space_after = Pt(14)
    run = summary.add_run(f"Jumlah: {len(questions)} soalan dengan paparan gambar. Nombor dan teks di bawah diambil terus daripada data portal semasa.")
    set_run_font(run, size=10.5, color=MUTED)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [1008, 8352]
    header_row = table.rows[0]
    repeat_header(header_row)
    prevent_row_split(header_row)
    headers = ["Nombor", "Soalan sebenar"]
    for index, text in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, MAROON)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5, bold=True, color=RGBColor(255, 255, 255))

    for row_index, (number, question) in enumerate(questions, start=1):
        row = table.add_row()
        prevent_row_split(row)
        if row_index % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        number_para = row.cells[0].paragraphs[0]
        number_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        number_para.paragraph_format.space_after = Pt(0)
        number_run = number_para.add_run(str(number))
        set_run_font(number_run, size=11, bold=True, color=RGBColor(138, 21, 56))

        question_para = row.cells[1].paragraphs[0]
        question_para.paragraph_format.space_after = Pt(0)
        question_para.paragraph_format.line_spacing = 1.15
        question_run = question_para.add_run(question)
        set_run_font(question_run, size=10.5, color=INK)

    set_table_geometry(table, widths)
    doc.core_properties.title = "Senarai Soalan Bergambar Sains Set 3"
    doc.core_properties.subject = "Portal PKSK"
    doc.core_properties.author = "Portal PKSK"
    doc.save(OUTPUT)


if __name__ == "__main__":
    question_rows = load_questions()
    if len(question_rows) != 31:
        raise RuntimeError(f"Dijangka 31 soalan bergambar, dijumpai {len(question_rows)}.")
    build_document(question_rows)
    print(OUTPUT)
